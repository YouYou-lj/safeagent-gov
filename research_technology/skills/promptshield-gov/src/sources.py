"""Adapters for user, web, document, RAG, and memory input sources."""

from __future__ import annotations

import re
from collections.abc import Iterable
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from safeagent_gov.contracts import SourceEnvelope, SourceType
from safeagent_gov.errors import UnsafePackageError

from .normalization import normalize_text

DEFAULT_TRUST = {
    SourceType.USER_INPUT: 0.45,
    SourceType.WEB_PAGE: 0.20,
    SourceType.UPLOADED_PDF: 0.30,
    SourceType.UPLOADED_DOC: 0.30,
    SourceType.RAG_RESULT: 0.55,
    SourceType.HISTORY_MEMORY: 0.50,
}
MAX_DOCUMENT_BYTES = 10 * 1024 * 1024
MAX_DOCUMENT_PAGES = 100


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._ignored_depth = 0
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.casefold() in {"script", "style", "noscript", "svg"}:
            self._ignored_depth += 1
        elif tag.casefold() in {"p", "br", "li", "div", "section", "article", "h1", "h2", "h3", "tr"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.casefold() in {"script", "style", "noscript", "svg"} and self._ignored_depth:
            self._ignored_depth -= 1
        elif tag.casefold() in {"p", "li", "div", "section", "article", "h1", "h2", "h3", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth:
            self.parts.append(data)

    def text(self) -> str:
        return " ".join(self.parts)


def _slug(value: str) -> str:
    compact = re.sub(r"[^A-Za-z0-9._:-]+", "-", value).strip("-")
    return compact[:80] or "source"


def create_source(
    content: str,
    source_type: SourceType | str,
    *,
    origin: str,
    trust_score: float | None = None,
    session_id: str | None = None,
    parent_source_id: str | None = None,
    metadata: dict[str, Any] | None = None,
    decode_html_entities: bool = False,
) -> SourceEnvelope:
    source_type = SourceType(source_type)
    normalized = normalize_text(content, decode_html_entities=decode_html_entities)
    source_id = f"{source_type.value}:{_slug(origin)}:{normalized.content_hash[:16]}"
    return SourceEnvelope(
        source_id=source_id,
        source_type=source_type,
        origin=origin,
        trust_score=DEFAULT_TRUST[source_type] if trust_score is None else trust_score,
        content=content,
        content_hash=normalized.content_hash,
        normalized_content=normalized.text,
        normalized_hash=normalized.normalized_hash,
        normalization_flags=list(normalized.flags),
        session_id=session_id,
        parent_source_id=parent_source_id,
        metadata=metadata or {},
    )


def adapt_user_input(
    text: str,
    *,
    user_id: str = "anonymous",
    session_id: str | None = None,
    metadata: dict[str, Any] | None = None,
) -> SourceEnvelope:
    return create_source(
        text,
        SourceType.USER_INPUT,
        origin=f"user:{user_id}",
        session_id=session_id,
        metadata={"user_id": user_id, **(metadata or {})},
    )


def adapt_web_content(
    html_content: str,
    *,
    url: str,
    session_id: str | None = None,
    trust_score: float | None = None,
    metadata: dict[str, Any] | None = None,
) -> SourceEnvelope:
    parsed_url = urlparse(url)
    if parsed_url.scheme not in {"http", "https"} or not parsed_url.hostname or parsed_url.username or parsed_url.password:
        raise ValueError("web source requires a valid HTTP/HTTPS URL without embedded credentials")
    parser = _VisibleTextParser()
    parser.feed(html_content[:1_000_000])
    visible_text = parser.text()
    return create_source(
        visible_text,
        SourceType.WEB_PAGE,
        origin=url,
        trust_score=trust_score,
        session_id=session_id,
        metadata={"hostname": parsed_url.hostname, "raw_html_hash_only": True, **(metadata or {})},
        decode_html_entities=True,
    )


def _extract_document(path: Path, ocr_text: str | None) -> tuple[str, SourceType, dict[str, Any]]:
    suffix = path.suffix.casefold()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        if len(reader.pages) > MAX_DOCUMENT_PAGES:
            raise UnsafePackageError(f"PDF page count exceeds {MAX_DOCUMENT_PAGES}")
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        method = "pdf_text"
        if not text.strip() and ocr_text:
            text, method = ocr_text, "ocr_fallback"
        return text, SourceType.UPLOADED_PDF, {"page_count": len(reader.pages), "extraction_method": method}
    if suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        text = "\n".join(paragraphs + table_cells)
        method = "docx_text"
        if not text.strip() and ocr_text:
            text, method = ocr_text, "ocr_fallback"
        return text, SourceType.UPLOADED_DOC, {"paragraph_count": len(paragraphs), "extraction_method": method}
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace"), SourceType.UPLOADED_DOC, {"extraction_method": "plain_text"}
    raise ValueError("supported document types are .pdf, .docx, .txt, and .md")


def adapt_document(
    path: str | Path,
    *,
    session_id: str | None = None,
    trust_score: float | None = None,
    ocr_text: str | None = None,
    metadata: dict[str, Any] | None = None,
) -> SourceEnvelope:
    document_path = Path(path).resolve()
    if not document_path.is_file():
        raise FileNotFoundError(f"document not found: {document_path}")
    size = document_path.stat().st_size
    if size > MAX_DOCUMENT_BYTES:
        raise UnsafePackageError(f"document size exceeds {MAX_DOCUMENT_BYTES} bytes")
    text, source_type, extracted = _extract_document(document_path, ocr_text)
    return create_source(
        text,
        source_type,
        origin=f"file:{document_path.name}",
        trust_score=trust_score,
        session_id=session_id,
        metadata={"filename": document_path.name, "size_bytes": size, **extracted, **(metadata or {})},
    )


def adapt_text_source(
    text: str,
    source_type: SourceType | str,
    *,
    origin: str,
    session_id: str | None = None,
    trust_score: float | None = None,
    metadata: dict[str, Any] | None = None,
) -> SourceEnvelope:
    return create_source(
        text,
        source_type,
        origin=origin,
        trust_score=trust_score,
        session_id=session_id,
        metadata=metadata,
    )


def adapt_rag_results(
    records: Iterable[dict[str, Any]],
    *,
    query_id: str,
    session_id: str | None = None,
) -> list[SourceEnvelope]:
    sources: list[SourceEnvelope] = []
    for rank, record in enumerate(records, 1):
        content = str(record.get("content", ""))
        origin = str(record.get("uri") or record.get("document_id") or f"rag:{query_id}:{rank}")
        score = float(record.get("retrieval_score", 0.0))
        trust_score = max(0.0, min(1.0, float(record.get("trust_score", DEFAULT_TRUST[SourceType.RAG_RESULT]))))
        sources.append(
            create_source(
                content,
                SourceType.RAG_RESULT,
                origin=origin,
                trust_score=trust_score,
                session_id=session_id,
                parent_source_id=str(record.get("document_id")) if record.get("document_id") else None,
                metadata={"query_id": query_id, "rank": rank, "retrieval_score": score, "citation": record.get("citation")},
            )
        )
    return sources


def adapt_memory_records(
    records: Iterable[dict[str, Any]],
    *,
    session_id: str,
) -> list[SourceEnvelope]:
    sources: list[SourceEnvelope] = []
    for turn, record in enumerate(records):
        role = str(record.get("role", "unknown"))
        memory_id = str(record.get("memory_id") or f"turn-{turn}")
        sources.append(
            create_source(
                str(record.get("content", "")),
                SourceType.HISTORY_MEMORY,
                origin=f"memory:{memory_id}",
                trust_score=float(record.get("trust_score", DEFAULT_TRUST[SourceType.HISTORY_MEMORY])),
                session_id=session_id,
                parent_source_id=str(record.get("parent_source_id")) if record.get("parent_source_id") else None,
                metadata={"turn": turn, "role": role, "memory_id": memory_id},
            )
        )
    return sources
